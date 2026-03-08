"""Replace the neuroradiology section with content from pediatric_neuroradiology.docx."""
import json, re, sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

doc = Document('C:/Users/dylan/child-neuro-handbook/pediatric_neuroradiology.docx')

def slugify(text, max_len=60):
    return re.sub(r'[^a-z0-9]+', '-', text.lower()).strip('-')[:max_len]

def get_hyperlinks_map(doc):
    """Build a map of relationship IDs to URLs."""
    rels = {}
    for rid, rel in doc.part.rels.items():
        if rel.target_ref and rel.target_ref.startswith('http'):
            rels[rid] = rel.target_ref
    return rels

rels_map = get_hyperlinks_map(doc)

def runs_to_html(paragraph, rels_map):
    """Convert paragraph runs to HTML, preserving bold/italic/hyperlinks."""
    parts = []
    for child in paragraph._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'hyperlink':
            rid = child.get(qn('r:id'))
            url = rels_map.get(rid, '') if rid else ''
            link_text = ''
            for node in child.iter():
                ns_tag = node.tag.split('}')[-1] if '}' in node.tag else node.tag
                if ns_tag == 't' and node.text:
                    link_text += node.text
            link_text = link_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if url:
                parts.append(f'<a href="{url}" target="_blank" rel="noopener noreferrer">{link_text}</a>')
            else:
                parts.append(link_text)
        elif tag == 'r':
            # Regular run
            from docx.text.run import Run
            run = Run(child, paragraph)
            t = run.text
            if not t:
                continue
            t = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if run.bold:
                t = f'<strong>{t}</strong>'
            if run.italic:
                t = f'<em>{t}</em>'
            parts.append(t)
    return ''.join(parts)

def cell_html(cell, rels_map):
    """Extract cell content preserving bold, italic, hyperlinks, and line breaks."""
    para_parts = []
    for p in cell.paragraphs:
        html = runs_to_html(p, rels_map)
        if html:
            para_parts.append(html)
    return '<br>'.join(para_parts)

def table_to_html(table, rels_map):
    """Convert a docx table to HTML."""
    rows = table.rows
    if not rows:
        return ''

    # Check for single-cell callout box
    first_cells = rows[0].cells
    unique_ids = set(id(c._tc) for c in first_cells)
    if len(unique_ids) == 1 and len(rows) == 1:
        content = cell_html(first_cells[0], rels_map)
        if len(first_cells[0].text.strip()) > 60:
            return f'<div class="callout-box" style="background:#f0f9ff;border-left:4px solid #3b82f6;padding:12px 16px;border-radius:6px;margin:12px 0;font-size:14px;line-height:1.6">{content}</div>'

    html = '<div class="table-wrap"><table>\n<thead>\n<tr>'
    # Header row — deduplicate merged cells
    seen_ids = set()
    header_count = 0
    for cell in first_cells:
        cid = id(cell._tc)
        if cid in seen_ids:
            continue
        seen_ids.add(cid)
        colspan = sum(1 for c in first_cells if id(c._tc) == cid)
        colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ''
        html += f'<th{colspan_attr}>{cell_html(cell, rels_map)}</th>'
        header_count += 1
    html += '</tr>\n</thead>\n<tbody>\n'

    for row in rows[1:]:
        html += '<tr>'
        seen_ids = set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen_ids:
                continue
            seen_ids.add(cid)
            colspan = sum(1 for c in row.cells if id(c._tc) == cid)
            colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ''
            html += f'<td{colspan_attr}>{cell_html(cell, rels_map)}</td>'
        html += '</tr>\n'

    html += '</tbody>\n</table></div>'
    return html

# ── Build HTML from document body elements ──────────────────────────────────

body = doc.element.body
elements = []
para_idx = 0
tbl_idx = 0
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        if para_idx < len(doc.paragraphs):
            elements.append(('para', doc.paragraphs[para_idx]))
            para_idx += 1
    elif tag == 'tbl':
        if tbl_idx < len(doc.tables):
            elements.append(('table', doc.tables[tbl_idx]))
            tbl_idx += 1

skip_texts = {
    'Pediatric Neuroimaging',
    'Child Neurology Residency Manual  |  PGY-3/4 Reference',
    'Child Neurology Residency Manual  |  Pediatric Neuroimaging  |  PGY-3/4',
}

html_parts = []
used_ids = set()

for etype, elem in elements:
    if etype == 'para':
        text = elem.text.strip()
        if not text:
            continue
        if text in skip_texts:
            continue
        if re.match(r'^Table \d+\.', text):
            html_parts.append(f'<p><strong>{text}</strong></p>')
            continue

        style = elem.style.name if elem.style else ''
        if style == 'Heading 1':
            clean = re.sub(r'^\d+\.\s*', '', text)
            sid = slugify(clean)
            base = sid
            counter = 2
            while sid in used_ids:
                sid = f'{base}-{counter}'
                counter += 1
            used_ids.add(sid)
            html_parts.append(f'<h2 id="{sid}">{clean}</h2>')
        elif style == 'Heading 2':
            sid = slugify(text)
            base = sid
            counter = 2
            while sid in used_ids:
                sid = f'{base}-{counter}'
                counter += 1
            used_ids.add(sid)
            html_parts.append(f'<h3 id="{sid}">{text}</h3>')
        else:
            para_html = runs_to_html(elem, rels_map)
            if para_html:
                html_parts.append(f'<p>{para_html}</p>')
    elif etype == 'table':
        html_parts.append(table_to_html(elem, rels_map))

new_html = '\n'.join(html_parts)

# Count links
link_count = new_html.count('<a href=')
table_count = new_html.count('<table>')
callout_count = new_html.count('callout-box')
print(f'Extracted: {len(html_parts)} elements, {table_count} tables, {callout_count} callouts, {link_count} links')

# ── Replace in section JSON ─────────────────────────────────────────────────

fpath = 'src/data/neuroradiology.json'
with open(fpath, 'r', encoding='utf-8') as f:
    data = json.load(f)

# The neuroradiology section has a single h1 "Reading MRIs" currently
# We'll replace the entire html with an h1 header + new content
section_h1 = '<h1 id="neuroradiology" style="color:#475569">Neuroradiology</h1>'
data['html'] = section_h1 + '\n' + new_html

# ── Build new TOC ───────────────────────────────────────────────────────────

new_toc = [{'level': 1, 'text': 'Neuroradiology', 'id': 'neuroradiology'}]
for part in html_parts:
    m = re.match(r'<h([23]) id="([^"]+)">(.*?)</h[23]>', part)
    if m:
        level = int(m.group(1))
        new_toc.append({
            'level': level,
            'text': re.sub(r'<[^>]+>', '', m.group(3)),
            'id': m.group(2)
        })

data['toc'] = new_toc

print(f'TOC: {len(new_toc)} entries')
for t in new_toc:
    print(f'  L{t["level"]} {t["text"]}  #{t["id"]}')

# Save
with open(fpath, 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False)
print(f'Saved {fpath}')

# Update index.json
with open('src/data/index.json', 'r', encoding='utf-8') as f:
    index = json.load(f)
for entry in index:
    if entry['slug'] == 'neuroradiology':
        entry['tocCount'] = len(new_toc)
        break
with open('src/data/index.json', 'w', encoding='utf-8') as f:
    json.dump(index, f, ensure_ascii=False, indent=2)
print(f'Updated index.json tocCount to {len(new_toc)}')
