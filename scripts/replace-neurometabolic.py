"""Replace the Neuro-Metabolic Disorders section with content from neurometabolic_disorders.docx."""
import json, re, sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

# ── 1. Extract HTML from docx ──────────────────────────────────────────────

doc = Document('C:/Users/dylan/child-neuro-handbook/neurometabolic_disorders.docx')

def slugify(text, max_len=60):
    return re.sub(r'[^a-z0-9]+', '-', text.lower()).strip('-')[:max_len]

def cell_html(cell):
    """Extract cell content, preserving bold and line breaks."""
    parts = []
    for p in cell.paragraphs:
        runs = []
        for r in p.runs:
            t = r.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if not t:
                continue
            if r.bold:
                t = f'<strong>{t}</strong>'
            if r.italic:
                t = f'<em>{t}</em>'
            runs.append(t)
        if runs:
            parts.append(''.join(runs))
    return '<br>'.join(parts)

def table_to_html(table):
    """Convert a docx table to an HTML table."""
    rows = table.rows
    if not rows:
        return ''

    # Check if first row is a single merged cell (callout box)
    first_row_cells = rows[0].cells
    unique_texts = set(c.text.strip() for c in first_row_cells)
    if len(rows) <= 2 and len(unique_texts) == 1 and len(first_row_cells[0].text.strip()) > 80:
        # This is a callout/info box, not a data table
        content = cell_html(first_row_cells[0])
        html = f'<div class="callout-box" style="background:#f0f9ff;border-left:4px solid #3b82f6;padding:12px 16px;border-radius:6px;margin:12px 0;font-size:14px;line-height:1.6">{content}</div>'
        return html

    html = '<div class="table-wrap"><table>\n<thead>\n<tr>'
    # Header row
    seen = set()
    for cell in first_row_cells:
        txt = cell_html(cell)
        if txt in seen:
            continue
        seen.add(txt)
        html += f'<th>{txt}</th>'
    html += '</tr>\n</thead>\n<tbody>\n'

    for row in rows[1:]:
        html += '<tr>'
        seen = set()
        for i, cell in enumerate(row.cells):
            # Skip merged cells (same object as previous)
            cell_id = id(cell._tc)
            if cell_id in seen:
                continue
            seen.add(cell_id)
            # Calculate colspan
            colspan = sum(1 for c in row.cells if id(c._tc) == cell_id)
            colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ''
            html += f'<td{colspan_attr}>{cell_html(cell)}</td>'
        html += '</tr>\n'

    html += '</tbody>\n</table></div>'
    return html

# Build HTML from paragraphs and tables
html_parts = []
used_ids = set()

# Track which tables we've already inserted (by index)
table_idx = 0
# Map paragraph indices to table insertions
# We need to figure out where tables appear in the document flow
# Tables in docx appear between paragraphs. We'll track by looking at the XML body children.

body = doc.element.body
elements = []  # list of ('para', paragraph) or ('table', table)

para_idx = 0
tbl_idx = 0
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        if para_idx < len(doc.paragraphs):
            elements.append(('para', doc.paragraphs[para_idx]))
            para_idx += 1
    elif tag == 'tbl':
        if tbl_idx < len(doc.tables):
            elements.append(('table', doc.tables[tbl_idx]))
            tbl_idx += 1

# Skip the title paragraphs (first two: "Neurometabolic Disorders" and subtitle)
skip_texts = {'Neurometabolic Disorders', 'Child Neurology Residency Manual  |  PGY-3/4 Reference',
              'Child Neurology Residency Manual  |  Neurometabolic Disorders  |  PGY-3/4'}

for etype, elem in elements:
    if etype == 'para':
        text = elem.text.strip()
        if not text:
            continue
        if text in skip_texts:
            continue
        # Table caption lines (e.g. "Table 1. ...")
        if re.match(r'^Table \d+\.', text):
            # Add as a caption paragraph
            html_parts.append(f'<p><strong>{text}</strong></p>')
            continue

        style = elem.style.name if elem.style else ''
        if style == 'Heading 1':
            # Remove leading number: "1. Conceptual Framework" -> "Conceptual Framework"
            clean = re.sub(r'^\d+\.\s*', '', text)
            sid = slugify(clean)
            base = sid
            counter = 2
            while sid in used_ids:
                sid = f'{base}-{counter}'
                counter += 1
            used_ids.add(sid)
            html_parts.append(f'<h2 id="{sid}">{clean}</h2>')
        elif style == 'Heading 2':
            sid = slugify(text)
            base = sid
            counter = 2
            while sid in used_ids:
                sid = f'{base}-{counter}'
                counter += 1
            used_ids.add(sid)
            html_parts.append(f'<h3 id="{sid}">{text}</h3>')
        else:
            # Regular paragraph - preserve bold/italic from runs
            parts = []
            for r in elem.runs:
                t = r.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                if not t:
                    continue
                if r.bold:
                    t = f'<strong>{t}</strong>'
                if r.italic:
                    t = f'<em>{t}</em>'
                parts.append(t)
            if parts:
                html_parts.append(f'<p>{"".join(parts)}</p>')
    elif etype == 'table':
        html_parts.append(table_to_html(elem))

new_section_html = '\n'.join(html_parts)
print(f'Extracted {len(html_parts)} HTML elements from docx')

# ── 2. Replace section in the JSON ──────────────────────────────────────────

fpath = 'src/data/neurogenetics-and-neurometabolics.json'
with open(fpath, 'r', encoding='utf-8') as f:
    data = json.load(f)

html = data['html']

# Find the h1 for Neuro-Metabolic and the next h1 after it
h1_positions = [(m.start(), m.group(1)) for m in re.finditer(r'<h1[^>]*id="([^"]+)"[^>]*>', html)]

start_pos = None
end_pos = None
for i, (pos, hid) in enumerate(h1_positions):
    if hid == 'neuro-metabolic-disorders':
        start_pos = pos
        if i + 1 < len(h1_positions):
            end_pos = h1_positions[i + 1][0]
        else:
            end_pos = len(html)
        break

if start_pos is None:
    print('ERROR: Could not find neuro-metabolic-disorders section')
    sys.exit(1)

old_section = html[start_pos:end_pos]
print(f'Old section: {len(old_section)} chars')

# Build the replacement: keep the h1 wrapper but replace content
new_h1 = '<h1 id="neuro-metabolic-disorders" style="color:#4f46e5">Neuro\u2011Metabolic Disorders: General Approach</h1>'
replacement = new_h1 + '\n' + new_section_html

new_html = html[:start_pos] + replacement + html[end_pos:]
data['html'] = new_html

# ── 3. Update TOC ──────────────────────────────────────────────────────────

# Remove old TOC entries for this section (indices 19-27)
old_toc = data['toc']
new_toc = [t for t in old_toc if not (
    t['id'] in {'metabolic-conceptual-split', 'metabolic-clinical-questions',
                'basic-metabolic-workup', 'treatable-first', 'mitochondrial-disorders',
                'mito-syndromes', 'leukodystrophies', 'test-now-vs-defer',
                'neuro-metabolic-disorders'}
)]

# Find insertion point: after neurocutaneous entries, before gdd-epilepsy-workup
insert_idx = None
for i, t in enumerate(new_toc):
    if t['id'] == 'gdd-epilepsy-workup':
        insert_idx = i
        break

# Build new TOC entries from the extracted headings
new_toc_entries = [{'level': 1, 'text': 'Neuro\u2011Metabolic Disorders', 'id': 'neuro-metabolic-disorders'}]
for part in html_parts:
    m = re.match(r'<h([23]) id="([^"]+)">(.*?)</h[23]>', part)
    if m:
        level = int(m.group(1))  # h2->2, h3->3
        new_toc_entries.append({
            'level': level,
            'text': re.sub(r'<[^>]+>', '', m.group(3)),
            'id': m.group(2)
        })

print(f'New TOC entries: {len(new_toc_entries)}')
for e in new_toc_entries:
    print(f'  L{e["level"]} {e["text"]}  #{e["id"]}')

if insert_idx is not None:
    new_toc = new_toc[:insert_idx] + new_toc_entries + new_toc[insert_idx:]
else:
    new_toc.extend(new_toc_entries)

data['toc'] = new_toc

# Update tocCount in index.json
data['tocCount'] = len(new_toc)

# ── 4. Save ─────────────────────────────────────────────────────────────────

with open(fpath, 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False)
print(f'Saved updated section JSON')

# Update index.json tocCount
with open('src/data/index.json', 'r', encoding='utf-8') as f:
    index = json.load(f)
for entry in index:
    if entry['slug'] == 'neurogenetics-and-neurometabolics':
        entry['tocCount'] = len(new_toc)
        break
with open('src/data/index.json', 'w', encoding='utf-8') as f:
    json.dump(index, f, ensure_ascii=False, indent=2)
print(f'Updated index.json tocCount to {len(new_toc)}')

print('\nDone! Run fix-heading-ids.py to rebuild search index.')
