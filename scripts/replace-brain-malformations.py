"""Replace Malformations of Brain Development + Skull/Posterior Fossa sections
in neurogenetics-and-neurometabolics.json with content from brain_development_malformations.docx."""
import json, re, sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

doc = Document('C:/Users/dylan/child-neuro-handbook/brain_development_malformations.docx')

def slugify(text, max_len=60):
    return re.sub(r'[^a-z0-9]+', '-', text.lower()).strip('-')[:max_len]

def runs_to_html(paragraph):
    """Convert paragraph runs to HTML, preserving bold/italic."""
    parts = []
    for child in paragraph._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            from docx.text.run import Run
            run = Run(child, paragraph)
            t = run.text
            if not t:
                continue
            t = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if run.bold:
                t = f'<strong>{t}</strong>'
            if run.italic:
                t = f'<em>{t}</em>'
            parts.append(t)
    return ''.join(parts)

def cell_html(cell):
    """Extract cell content preserving bold, italic, and line breaks."""
    para_parts = []
    for p in cell.paragraphs:
        html = runs_to_html(p)
        if html:
            para_parts.append(html)
    return '<br>'.join(para_parts)

def table_to_html(table):
    """Convert a docx table to HTML."""
    rows = table.rows
    if not rows:
        return ''

    # Check for single-cell callout box
    first_cells = rows[0].cells
    unique_ids = set(id(c._tc) for c in first_cells)
    if len(unique_ids) == 1 and len(rows) == 1:
        content = cell_html(first_cells[0])
        if len(first_cells[0].text.strip()) > 60:
            return (f'<div class="callout-box" style="background:#f0f9ff;border-left:4px solid #3b82f6;'
                    f'padding:12px 16px;border-radius:6px;margin:12px 0;font-size:14px;line-height:1.6">'
                    f'{content}</div>')

    html = '<div class="table-wrap"><table>\n<thead>\n<tr>'
    seen_ids = set()
    for cell in first_cells:
        cid = id(cell._tc)
        if cid in seen_ids:
            continue
        seen_ids.add(cid)
        colspan = sum(1 for c in first_cells if id(c._tc) == cid)
        colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ''
        html += f'<th{colspan_attr}>{cell_html(cell)}</th>'
    html += '</tr>\n</thead>\n<tbody>\n'

    for row in rows[1:]:
        html += '<tr>'
        seen_ids = set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen_ids:
                continue
            seen_ids.add(cid)
            colspan = sum(1 for c in row.cells if id(c._tc) == cid)
            colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ''
            html += f'<td{colspan_attr}>{cell_html(cell)}</td>'
        html += '</tr>\n'

    html += '</tbody>\n</table></div>'
    return html

# ── Build HTML from document body elements ──────────────────────────────────

body = doc.element.body
elements = []
para_idx = 0
tbl_idx = 0
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        if para_idx < len(doc.paragraphs):
            elements.append(('para', doc.paragraphs[para_idx]))
            para_idx += 1
    elif tag == 'tbl':
        if tbl_idx < len(doc.tables):
            elements.append(('table', doc.tables[tbl_idx]))
            tbl_idx += 1

skip_texts = {
    'Brain Development and Malformations',
    'Child Neurology Residency Manual  |  Neurogenetics Section  |  PGY-3/4 Reference',
    'Child Neurology Residency Manual  |  Brain Development and Malformations  |  PGY-3/4',
}

html_parts = []
used_ids = set()

for etype, elem in elements:
    if etype == 'para':
        text = elem.text.strip()
        if not text:
            continue
        if text in skip_texts:
            continue
        if re.match(r'^Table \d+\.', text):
            html_parts.append(f'<p><strong>{text}</strong></p>')
            continue

        style = elem.style.name if elem.style else ''
        if style == 'Heading 1':
            clean = re.sub(r'^\d+\.\s*', '', text)
            sid = slugify(clean)
            base = sid
            counter = 2
            while sid in used_ids:
                sid = f'{base}-{counter}'
                counter += 1
            used_ids.add(sid)
            html_parts.append(f'<h2 id="{sid}">{clean}</h2>')
        elif style == 'Heading 2':
            sid = slugify(text)
            base = sid
            counter = 2
            while sid in used_ids:
                sid = f'{base}-{counter}'
                counter += 1
            used_ids.add(sid)
            html_parts.append(f'<h3 id="{sid}">{text}</h3>')
        else:
            para_html = runs_to_html(elem)
            if para_html:
                html_parts.append(f'<p>{para_html}</p>')
    elif etype == 'table':
        html_parts.append(table_to_html(elem))

new_section_html = '\n'.join(html_parts)

table_count = new_section_html.count('<table>')
callout_count = new_section_html.count('callout-box')
print(f'Extracted: {len(html_parts)} elements, {table_count} tables, {callout_count} callouts')

# ── Replace in section JSON ─────────────────────────────────────────────────

fpath = 'src/data/neurogenetics-and-neurometabolics.json'
with open(fpath, 'r', encoding='utf-8') as f:
    data = json.load(f)

html = data['html']

# Find the boundary: from start (pos 0) to h1#neurocutaneous-disorders
neurocutaneous_pos = html.index('id="neurocutaneous-disorders"')
neurocutaneous_pos = html.rfind('<h1', 0, neurocutaneous_pos)

old_section = html[:neurocutaneous_pos]
print(f'Old section: {len(old_section)} chars')

# Build replacement with h1 header
h1_tag = '<h1 id="brain-development-and-malformations" style="color:#4f46e5">Brain Development and Malformations</h1>'
replacement = h1_tag + '\n' + new_section_html + '\n'

new_html = replacement + html[neurocutaneous_pos:]
data['html'] = new_html

# ── Update TOC ──────────────────────────────────────────────────────────────

# Remove old TOC entries for the replaced sections
old_ids_to_remove = {
    'malformations-of-brain-development', 'developmental-timing',
    'selected-brain-malformations', 'agenesis-corpus-callosum',
    'holoprosencephaly', 'optic-nerve-hypoplasia', 'hemimegalencephaly',
    'lissencephaly', 'skull-posterior-fossa', 'craniosynostosis',
    'chiari-deformation', 'dandy-walker',
}
toc = [t for t in data['toc'] if t['id'] not in old_ids_to_remove]

# Build new TOC entries
new_toc_entries = [{'level': 1, 'text': 'Brain Development and Malformations', 'id': 'brain-development-and-malformations'}]
for part in html_parts:
    m = re.match(r'<h([23]) id="([^"]+)">(.*?)</h[23]>', part)
    if m:
        level = int(m.group(1))
        new_toc_entries.append({
            'level': level,
            'text': re.sub(r'<[^>]+>', '', m.group(3)),
            'id': m.group(2)
        })

# Insert at the beginning (before neurocutaneous)
neurocutaneous_idx = next(i for i, t in enumerate(toc) if t['id'] == 'neurocutaneous-disorders')
toc = toc[:neurocutaneous_idx] + new_toc_entries + toc[neurocutaneous_idx:]
data['toc'] = toc

print(f'\nNew TOC entries: {len(new_toc_entries)}')
for e in new_toc_entries:
    print(f'  L{e["level"]} {e["text"]}  #{e["id"]}')

# Save
with open(fpath, 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False)
print(f'\nSaved {fpath}')

# Update index.json
with open('src/data/index.json', 'r', encoding='utf-8') as f:
    index = json.load(f)
for entry in index:
    if entry['slug'] == 'neurogenetics-and-neurometabolics':
        entry['tocCount'] = len(toc)
        break
with open('src/data/index.json', 'w', encoding='utf-8') as f:
    json.dump(index, f, ensure_ascii=False, indent=2)
print(f'Updated index.json tocCount to {len(toc)}')
